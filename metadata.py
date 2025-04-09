import os
import re
import glob
import json
import sys
from docx import Document
from pdfminer.high_level import extract_text
import pytesseract
from pdf2image import convert_from_path

sys.setrecursionlimit(10000)

def count_tokens(text):
    words = text.split()
    return len(words)


def has_selectable_text(pdf_path):
    try:
        text = extract_text(pdf_path)
        return len(text.strip()) > 0
    except Exception as e:
        print(f"Error checking selectable text in {pdf_path}: {e}")
        return False


def extract_text_from_pdf_with_ocr(pdf_path):
    try:
        images = convert_from_path(pdf_path)
        text = ""
        for image in images:
            text += pytesseract.image_to_string(image, lang="vie")
        return text
    except Exception as e:
        print(f"OCR error for {pdf_path}: {e}")
        return ""


def extract_text_from_pdf(pdf_path):
    try:
        if has_selectable_text(pdf_path):
            return extract_text(pdf_path)
        else:
            return extract_text_from_pdf_with_ocr(pdf_path)
    except RecursionError:
        print(f"Recursion error in PDF processing for {pdf_path}. Trying alternative method.")
        try:
            return extract_text_from_pdf_with_ocr(pdf_path)
        except Exception as e:
            print(f"Alternative method failed for {pdf_path}: {e}")
            return ""


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    paragraphs_text = [para.text for para in doc.paragraphs]
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            tables_text.append(" | ".join(row_text))
    all_text = "\n".join(paragraphs_text + tables_text)
    return all_text


def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError("Unsupported file type")

def extract_document_id(text):
    patterns = [
        r"Số:\s*(\d+/\d+/TT-\w+)",
        r"Số:\s*(\d+\.\d+\.TT\.[\w\.]+)",
        r"(TT\.\d+\.\d+\.[\w\.]+)",
        r"([A-Z]+\d+\.signed)",
        r"(\d+[_/]\d+[_/](?:TT|TTLT)[-_][\w]+)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1)

    return "Unknown"


def extract_title(text):
    lines = text.split("\n")

    for i, line in enumerate(lines):
        if "THÔNG TƯ" in line.upper():
            for j in range(i+1, min(i+5, len(lines))):
                if lines[j].strip() and "BỘ" not in lines[j] and "Căn cứ" not in lines[j]:
                    return lines[j].strip()
    
    return "Unknown"

def extract_issuing_agency(text):
    patterns = [
        r"(BỘ [^\n]+)",
        r"(CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM[\s\S]*?)\n"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    
    return "Unknown"


def extract_date(text):
    patterns = [
        r"Hà Nội, (ngày \d+ tháng \d+ năm \d+)",
        r"ngày (\d+ tháng \d+ năm \d+)",
        r"(\d+/\d+/\d+)",
        r"(\d+ tháng \d+ năm \d+)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    
    return "Unknown"


def split_into_segments(text):
    pattern = r"(Điều \d+\..*?)(?=Điều \d+\.|$)"
    segments = re.findall(pattern, text, flags=re.DOTALL)
    if not segments:
        pattern = r"(Chương \d+\..*?)(?=Chương \d+\.|$)"
        segments = re.findall(pattern, text, flags=re.DOTALL)
    if not segments and len(text) > 1000:
        segments = [text[i:i+1000] for i in range(0, len(text), 1000)]
    elif not segments:
        segments = [text]
    
    segments = [s.strip() for s in segments if s.strip()]
    token_limit = 256
    final_segments = []
    
    for segment in segments:
        title_match = re.match(r"(Điều \d+[\.:][^\n]*|Chương \d+[\.:][^\n]*)", segment)
        segment_title = title_match.group(0) if title_match else ""
        if count_tokens(segment) <= token_limit:
            final_segments.append(segment)
            continue
        
        words = segment.split()
        current_segment_words = []
        current_token_count = 0
        part_counter = 1
        
        title_token_count = count_tokens(segment_title) if segment_title else 0
        
        for word in words:
            if current_token_count + 1 <= token_limit:
                current_segment_words.append(word)
                current_token_count += 1
            else:
                current_segment = " ".join(current_segment_words)
                if segment_title and part_counter == 1:
                    final_segments.append(f"{segment_title}\n{current_segment}")
                elif segment_title:
                    final_segments.append(f"{segment_title} (phần {part_counter})\n{current_segment}")
                else:
                    final_segments.append(current_segment)
                
                current_segment_words = [word]
                current_token_count = 1
                part_counter += 1
        
        if current_segment_words:
            current_segment = " ".join(current_segment_words)
            if segment_title and part_counter == 1:
                final_segments.append(f"{segment_title}\n{current_segment}")
            elif segment_title:
                final_segments.append(f"{segment_title} (phần {part_counter})\n{current_segment}")
            else:
                final_segments.append(current_segment)
    
    return final_segments


directory = "/home/cong/workspace/chatbot-retrieval-based/Thongtu"
docx_files = glob.glob(os.path.join(directory, "*.docx"))
pdf_files = glob.glob(os.path.join(directory, "*.pdf"))
all_files = docx_files + pdf_files

segment_data = []
for file_path in all_files:
    try:
        print(f"Processing {file_path}...")
        text = extract_text(file_path)
        
        filename = os.path.basename(file_path)
        doc_id = extract_document_id(text)

        if doc_id == "Unknown":
            doc_id = os.path.splitext(filename)[0]
        
        title = extract_title(text)
        agency = extract_issuing_agency(text)
        date = extract_date(text)
        segments = split_into_segments(text)
        
        if not segments:
            print(f"Warning: No segments found in {file_path}")
            segments = [text]
            
        for i, segment in enumerate(segments):
            segment_data.append({
                "document_id": doc_id,
                "title": title,
                "issuing_agency": agency,
                "date": date,
                "segment_id": f"{doc_id}_{i}",
                "content": segment
            })
        print(f"Successfully processed {file_path} with {len(segments)} segments")
    except Exception as e:
        print(f"Error processing {file_path}: {e}")

with open("metadata.json", "w", encoding="utf-8") as f:
    json.dump(segment_data, f, ensure_ascii=False, indent=4)

print(f"Metadata saved to metadata.json with {len(segment_data)} segments from {len(all_files)} files")
print("\nĐể tạo vector embeddings, hãy chạy script embeddings.py tiếp theo.")